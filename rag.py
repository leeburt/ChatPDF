# -*- coding: utf-8 -*-
"""
@author:XuMing(xuming624@qq.com)
@description: 
"""
import argparse
import hashlib
import os
import re
from threading import Thread
from typing import Union, List
import pprint

import jieba
import torch
from loguru import logger
from peft import PeftModel

from cal_similarites import (
    MyEnsembleSimilarity as EnsembleSimilarity,
    MyBertSimilarity as BertSimilarity,
    MyBM25Similarity as BM25Similarity,
    SimilarityABC
)
from transformers import (
    AutoModel,
    AutoModelForCausalLM,
    AutoTokenizer,
    BloomForCausalLM,
    BloomTokenizerFast,
    LlamaTokenizer,
    LlamaForCausalLM,
    TextIteratorStreamer,
    GenerationConfig,
    AutoModelForSequenceClassification,
)
import openai

jieba.setLogLevel("ERROR")

MODEL_CLASSES = {
    "bloom": (BloomForCausalLM, BloomTokenizerFast),
    "chatglm": (AutoModel, AutoTokenizer),
    "llama": (LlamaForCausalLM, LlamaTokenizer),
    "baichuan": (AutoModelForCausalLM, AutoTokenizer),
    "auto": (AutoModelForCausalLM, AutoTokenizer),
}

PROMPT_TEMPLATE = """基于以下已知信息，专业的来回答用户的问题，用简体中文按照逻辑详细回答，同时标注对应的证据来源。回答时注意仔细阅读所有提供的资料也即是下方提供的已知内容，提取出与问题 {query_str} 相关的内容，并正确引用对应的已知内容的编号。
所有已知内容需要逐一查看，识别出与问题 {query_str} 相关的关键点,并正确引用对应的证据编号。同时要注意不同证据之间可能存在的矛盾，并确保输出的信息全面且准确
。

【已知内容】:
{context_str}

【问题】:
{query_str}

如果问题和已知内容不匹配，则不要解释直接输出:未找到有关 {query_str} 的结果。

"""



class SentenceSplitter:
    def __init__(self, chunk_size: int = 250, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[str]:
        if self._is_has_chinese(text):
            return self._split_chinese_text(text)
        else:
            return self._split_english_text(text)

    def _split_chinese_text(self, text: str) -> List[str]:
        sentence_endings = {'\n', '。', '！', '？', '；', '…'}  # 句末标点符号
        chunks, current_chunk = [], ''
        for word in jieba.cut(text):
            if len(current_chunk) + len(word) > self.chunk_size:
                chunks.append(current_chunk.strip())
                current_chunk = word
            else:
                current_chunk += word
            if word[-1] in sentence_endings and len(current_chunk) > self.chunk_size - self.chunk_overlap:
                chunks.append(current_chunk.strip())
                current_chunk = ''
        if current_chunk:
            chunks.append(current_chunk.strip())
        if self.chunk_overlap > 0 and len(chunks) > 1:
            chunks = self._handle_overlap(chunks)
        return chunks

    def _split_english_text(self, text: str) -> List[str]:
        # 使用正则表达式按句子分割英文文本
        sentences = re.split(r'(?<=[.!?])\s+', text.replace('\n', ' '))
        chunks = []
        current_chunk = ''
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= self.chunk_size:
                current_chunk += (' ' if current_chunk else '') + sentence
            else:
                if len(sentence) > self.chunk_size:
                    for i in range(0, len(sentence), self.chunk_size):
                        chunks.append(sentence[i:i + self.chunk_size])
                    current_chunk = ''
                else:
                    chunks.append(current_chunk)
                    current_chunk = sentence
        if current_chunk:  # Add the last chunk
            chunks.append(current_chunk)

        if self.chunk_overlap > 0 and len(chunks) > 1:
            chunks = self._handle_overlap(chunks)

        return chunks

    def _is_has_chinese(self, text: str) -> bool:
        # check if contains chinese characters
        if any("\u4e00" <= ch <= "\u9fff" for ch in text):
            return True
        else:
            return False

    def _handle_overlap(self, chunks: List[str]) -> List[str]:
        # 处理块间重叠
        overlapped_chunks = []
        for i in range(len(chunks) - 1):
            chunk = chunks[i] + ' ' + chunks[i + 1][:self.chunk_overlap]
            overlapped_chunks.append(chunk.strip())
        overlapped_chunks.append(chunks[-1])
        return overlapped_chunks
    
import os 
# 从环境变量中获取 OpenAI API 密钥
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY 环境变量未设置")

# 从环境变量中获取 OpenAI 模型名称
OPENAI_MODEL = os.getenv("OPENAI_MODEL")
if not OPENAI_MODEL:
    raise ValueError("OPENAI_MODEL 环境变量未设置")

# 从环境变量中获取 API 的基础 URL
BASE_URL = os.getenv("BASE_URL")
if not BASE_URL:
    raise ValueError("BASE_URL 环境变量未设置")

class Rag:
    def __init__(
            self,
            similarity_model: SimilarityABC = None,
            generate_model_type: str = "auto",
            generate_model_name_or_path: str = "Qwen/Qwen2-0.5B-Instruct",
            lora_model_name_or_path: str = None,
            corpus_files_path: str = "./corpus/",
            corpus_files: Union[str, List[str]] = None,
            save_corpus_emb_dir: str = "./corpus_embs/",
            device: str = None,
            int8: bool = False,
            int4: bool = False,
            chunk_size: int = 250,
            chunk_overlap: int = 0,
            rerank_model_name_or_path: str = None,
            enable_history: bool = False,
            num_expand_context_chunk: int = 10,
            similarity_top_k: int = 50,
            rerank_top_k: int = 3,
            openai_api_key: str =  OPENAI_API_KEY,
            openai_model: str = OPENAI_MODEL ,
            base_url :str = BASE_URL
    ):
        """
        Init RAG model.
        :param similarity_model: similarity model, default None, if set, will use it instead of EnsembleSimilarity
        :param generate_model_type: generate model type
        :param generate_model_name_or_path: generate model name or path
        :param lora_model_name_or_path: lora model name or path
        :param corpus_files: corpus files
        :param save_corpus_emb_dir: save corpus embeddings dir, default ./corpus_embs/
        :param device: device, default None, auto select gpu or cpu
        :param int8: use int8 quantization, default False
        :param int4: use int4 quantization, default False
        :param chunk_size: chunk size, default 250
        :param chunk_overlap: chunk overlap, default 0, can not set to > 0 if num_expand_context_chunk > 0
        :param rerank_model_name_or_path: rerank model name or path, default 'BAAI/bge-reranker-base'
        :param enable_history: enable history, default False
        :param num_expand_context_chunk: num expand context chunk, default 2, if set to 0, will not expand context chunk
        :param similarity_top_k: similarity_top_k, default 5, similarity model search k corpus chunks
        :param rerank_top_k: rerank_top_k, default 3, rerank model search k corpus chunks
        """
        if torch.cuda.is_available():
            default_device = torch.device(0)
        elif torch.backends.mps.is_available():
            default_device = torch.device('cpu')
        else:
            default_device = torch.device('cpu')
        self.device = device or default_device
        if num_expand_context_chunk > 0 and chunk_overlap > 0:
            logger.warning(f" 'num_expand_context_chunk' and 'chunk_overlap' cannot both be greater than zero. "
                           f" 'chunk_overlap' has been set to zero by default.")
            chunk_overlap = 0
        self.text_splitter = SentenceSplitter(chunk_size, chunk_overlap)
        if similarity_model is not None:
            self.sim_model = similarity_model
        else:
            m1 = BertSimilarity(save_path=corpus_files_path,model_name_or_path="shibing624/text2vec-base-multilingual", device=self.device)
            m2 = BM25Similarity(save_path=corpus_files_path)
            default_sim_model = EnsembleSimilarity(similarities=[m1, m2], weights=[0.5, 0.5], c=2)
            self.sim_model = default_sim_model
        self.gen_model, self.tokenizer = self._init_gen_model(
            generate_model_type,
            generate_model_name_or_path,
            peft_name=lora_model_name_or_path,
            int8=int8,
            int4=int4,
        )
        self.history = []
        corpus_files=self.get_corpus(corpus_files_path)
        print(f"所有文档数量为:{len(corpus_files)}")
        self.convert_docx_to_markdown(corpus_files_path)
        self.corpus_files = corpus_files
        if corpus_files:
            self.add_corpus(corpus_files)
        self.save_corpus_emb_dir = save_corpus_emb_dir
        if rerank_model_name_or_path is None:
            rerank_model_name_or_path = "BAAI/bge-reranker-base"
        if rerank_model_name_or_path:
            self.rerank_tokenizer = AutoTokenizer.from_pretrained(rerank_model_name_or_path)
            self.rerank_model = AutoModelForSequenceClassification.from_pretrained(rerank_model_name_or_path)
            self.rerank_model.to(self.device)
            self.rerank_model.eval()
        else:
            self.rerank_model = None
            self.rerank_tokenizer = None
        self.enable_history = enable_history
        self.similarity_top_k = similarity_top_k
        self.num_expand_context_chunk = num_expand_context_chunk
        self.rerank_top_k = rerank_top_k

        self.openai_api_key= openai_api_key
        self.openai_model = openai_model
        self.base_url = base_url

    def __str__(self):
        return f"Similarity model: {self.sim_model}, Generate model: {self.gen_model}"

    def _init_gen_model(
            self,
            gen_model_type: str,
            gen_model_name_or_path: str,
            peft_name: str = None,
            int8: bool = False,
            int4: bool = False,
    ):
        """Init generate model."""
        if int8 or int4:
            device_map = None
        else:
            device_map = "auto"
        model_class, tokenizer_class = MODEL_CLASSES[gen_model_type]
        tokenizer = tokenizer_class.from_pretrained(gen_model_name_or_path, trust_remote_code=True)
        model = model_class.from_pretrained(
            gen_model_name_or_path,
            load_in_8bit=int8 if gen_model_type not in ['baichuan', 'chatglm'] else False,
            load_in_4bit=int4 if gen_model_type not in ['baichuan', 'chatglm'] else False,
            torch_dtype="auto",
            device_map=device_map,
            trust_remote_code=True,
        )
        if self.device == torch.device('cpu'):
            model.float()
        if gen_model_type in ['baichuan', 'chatglm']:
            if int4:
                model = model.quantize(4).cuda()
            elif int8:
                model = model.quantize(8).cuda()
        try:
            model.generation_config = GenerationConfig.from_pretrained(gen_model_name_or_path, trust_remote_code=True)
        except Exception as e:
            logger.warning(f"Failed to load generation config from {gen_model_name_or_path}, {e}")
        if peft_name:
            model = PeftModel.from_pretrained(
                model,
                peft_name,
                torch_dtype="auto",
            )
            logger.info(f"Loaded peft model from {peft_name}")
        model.eval()
        return model, tokenizer

    def _get_chat_input(self):
        messages = []
        for conv in self.history:
            if conv and len(conv) > 0 and conv[0]:
                messages.append({'role': 'user', 'content': conv[0]})
            if conv and len(conv) > 1 and conv[1]:
                messages.append({'role': 'assistant', 'content': conv[1]})
        input_ids = self.tokenizer.apply_chat_template(
            conversation=messages,
            tokenize=True,
            add_generation_prompt=True,
            return_tensors='pt'
        )
        return input_ids.to(self.gen_model.device)

    @torch.inference_mode()
    def stream_generate_answer(
            self,
            max_new_tokens=512,
            temperature=0.7,
            repetition_penalty=1.0,
            context_len=2048
    ):
        streamer = TextIteratorStreamer(self.tokenizer, timeout=60.0, skip_prompt=True, skip_special_tokens=True)
        input_ids = self._get_chat_input()
        max_src_len = context_len - max_new_tokens - 8
        input_ids = input_ids[-max_src_len:]
        generation_kwargs = dict(
            input_ids=input_ids,
            max_new_tokens=max_new_tokens,
            temperature=temperature,
            do_sample=True,
            repetition_penalty=repetition_penalty,
            streamer=streamer,
        )
        thread = Thread(target=self.gen_model.generate, kwargs=generation_kwargs)
        thread.start()

        yield from streamer

    def get_corpus(self,file_path):
        if os.path.isdir(file_path):
            ## 递归查找所有的文档
            all_files = []

            for root, dirs, files in os.walk(file_path):
                for file in files:
                    if file.endswith('.docx'):
                        all_files.append(os.path.join(root, file))
        else:
            all_files = [file_path]
        return all_files
    
    def convert_docx_to_markdown(self,docx_path):
        from docx import Document
        def extract_text_from_docx(file_path):
            """
            从 .docx 文件中提取文本内容。
            """
            try:
                full_text = []
                doc = Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
            except:
                logger.error(f"Failed to load docx file: {file_path}")
            return '\n'.join(full_text)
        """
        递归遍历指定路径下的所有 .docx 文件，并将内容写入 Markdown 文件。
        """
        markdown_path=os.path.join(docx_path, 'all_corpus.md')
        # if os.path.exists(markdown_path):
        #     logger.debug(f"{markdown_path} exists, skip!!!")
        #     return
        
        with open(markdown_path, 'w', encoding='utf-8') as md_file:
            file_list=self.get_corpus(docx_path)
            # print(markdown_path,file_list)
            for file in file_list:
                file_name = os.path.basename(file)  # 去掉 .docx 后缀
                md_file.write(f"# {file_name}\n\n")
                # 提取 .docx 文件内容
                docx_content = extract_text_from_docx(file)
                # 将内容写入 Markdown 文件
                md_file.write(f"{docx_content}\n\n")
        logger.debug(f"write all corpus to {markdown_path}")
        return 

    def add_corpus(self, files: Union[str, List[str]]):
        """Load document files."""
        if isinstance(files, str):
            files = [files]
        # print("corpus:",files)
        all_chunks=[]
        all_doc_name=[]
        for doc_file in files:
            try:
                if doc_file.endswith('.pdf'):
                    corpus = self.extract_text_from_pdf(doc_file)
                elif doc_file.endswith('.docx'):
                    corpus = self.extract_text_from_docx(doc_file)
                elif doc_file.endswith('.md'):
                    corpus = self.extract_text_from_markdown(doc_file)
                else:
                    corpus = self.extract_text_from_txt(doc_file)
            except:
                print(f"文件读取失败:{doc_file}")
            full_text = '\n'.join(corpus)
            chunks = self.text_splitter.split_text(full_text)
            ##获取文件名称
            doc_name=[os.path.basename(doc_file)]*len(chunks)
            all_chunks.extend(chunks)
            all_doc_name.extend(doc_name)
        self.sim_model.add_corpus(all_chunks,all_doc_name)
        self.corpus_files = files
        # logger.debug(f"corpus size: {len(self.sim_model.corpus)}, top3: "
        #              f"{list(self.sim_model.corpus.values())[:3]}")
    @staticmethod
    def get_file_hash(fpaths):
        hasher = hashlib.md5()
        target_file_data = bytes()
        if isinstance(fpaths, str):
            fpaths = [fpaths]
        for fpath in fpaths:
            with open(fpath, 'rb') as file:
                chunk = file.read(1024 * 1024)  # read only first 1MB
                hasher.update(chunk)
                target_file_data += chunk

        hash_name = hasher.hexdigest()[:32]
        return hash_name

    @staticmethod
    def extract_text_from_pdf(file_path: str):
        """Extract text content from a PDF file."""
        import PyPDF2
        contents = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text().strip()
                raw_text = [text.strip() for text in page_text.splitlines() if text.strip()]
                new_text = ''
                for text in raw_text:
                    new_text += text
                    if text[-1] in ['.', '!', '?', '。', '！', '？', '…', ';', '；', ':', '：', '”', '’', '）', '】', '》', '」',
                                    '』', '〕', '〉', '》', '〗', '〞', '〟', '»', '"', "'", ')', ']', '}']:
                        contents.append(new_text)
                        new_text = ''
                if new_text:
                    contents.append(new_text)
        return contents

    @staticmethod
    def extract_text_from_txt(file_path: str):
        """Extract text content from a TXT file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            contents = [text.strip() for text in f.readlines() if text.strip()]
        return contents

    @staticmethod
    def extract_text_from_docx(file_path: str):
        """Extract text content from a DOCX file."""
        import docx
        document = docx.Document(file_path)
        contents = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return contents

    @staticmethod
    def extract_text_from_markdown(file_path: str):
        """Extract text content from a Markdown file."""
        import markdown
        from bs4 import BeautifulSoup
        with open(file_path, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        html = markdown.markdown(markdown_text)
        soup = BeautifulSoup(html, 'html.parser')
        contents = [text.strip() for text in soup.get_text().splitlines() if text.strip()]
        return contents


    def openai_stream_generate_answer(self, prompt: str, max_tokens: int = 512, temperature: float = 0.7):
        from openai import OpenAI
        client = OpenAI(api_key=self.openai_api_key,base_url=self.base_url)

        response = client.chat.completions.create(
            model=self.openai_model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens,
            temperature=temperature,
            stream=True
        )

        for chunk in response:
            if chunk.choices[0].delta.content:
                yield chunk.choices[0].delta.content
    @staticmethod
    def _add_source_numbers(lst):
        """Add source numbers to a list of strings."""
        return [f'[{idx + 1}]\t "{item}"' for idx, item in enumerate(lst)]

    def _get_reranker_score(self, query: str, reference_results: List[str]):
        """Get reranker score."""
        pairs = []
        for reference in reference_results:
            pairs.append([query, reference])
        with torch.no_grad():
            inputs = self.rerank_tokenizer(pairs, padding=True, truncation=True, return_tensors='pt', max_length=512)
            inputs_on_device = {k: v.to(self.rerank_model.device) for k, v in inputs.items()}
            scores = self.rerank_model(**inputs_on_device, return_dict=True).logits.view(-1, ).float()

        return scores

    def get_reference_results(self, query: str):
        """
        Get reference results.
            1. Similarity model get similar chunks
            2. Rerank similar chunks
            3. Expand reference context chunk
        :param query:
        :return:
        """
        reference_results = []
        sim_contents = self.sim_model.most_similar(query, topn=self.similarity_top_k)
        # Get reference results from corpus
        hit_chunk_dict = dict()
        hit_chunk_info_dict=dict()
        for c in sim_contents:
            for id_score_dict in c:
                corpus_id = id_score_dict['corpus_id']
                hit_chunk = id_score_dict["corpus_doc"]
                doc_name = id_score_dict["doc_name"]
                score = id_score_dict["score"]
                reference_results.append(hit_chunk)
                hit_chunk_dict[corpus_id] = hit_chunk
                hit_chunk_info_dict[corpus_id]={
                    "corpus_id":corpus_id,
                    "hit_chunk":hit_chunk,
                    "doc_name":doc_name,
                    "score":score,
                    "rerank_score":-99999.0,
                    "is_rerank":False
                }

        logger.debug(f"reference_results: {len(sim_contents)},self.rerank_model:{self.rerank_model}")

        if reference_results:
            if self.rerank_model is not None:
                # Rerank reference results
                rerank_scores = self._get_reranker_score(query, reference_results)
                # logger.debug(f"rerank_scores: {rerank_scores}")
                # Get rerank top k chunks
                reference_results = [(reference,score) for reference, score in sorted(
                    zip(reference_results, rerank_scores), key=lambda x: x[1], reverse=True)][:self.rerank_top_k]
                ##获取reference_results第一列的值
                # reference_results_chunks = [reference for reference, score in reference_results]

                for corpus_id, hit_chunk_info in hit_chunk_info_dict.items():
                    hit_chunk=hit_chunk_info['hit_chunk']
                    for reference, rerank_score in reference_results:
                        if reference == hit_chunk:
                            ##把rerank_score 的tensor格式转为cpu的数据类型,正常的float
                            rerank_score = rerank_score.to(torch.float32)
                            hit_chunk_info_dict[corpus_id]["rerank_score"] = rerank_score.item()
                            hit_chunk_info_dict[corpus_id]["is_rerank"] = True
                
                # hit_chunk_dict = {corpus_id: hit_chunk for corpus_id, hit_chunk in hit_chunk_dict.items() if
                #                   hit_chunk in reference_results_chunks}
                
                ##按照hit_chunk_info_dict中的rerank_score排序，返回排序后的hit_chunk_info_dict
                hit_chunk_info_dict=dict(sorted(hit_chunk_info_dict.items(), key=lambda x: x[1]["rerank_score"], reverse=True))
        
            # Expand reference context chunk
            new_hit_chunk_info_dict=hit_chunk_info_dict
            print("hit_chunk_info_dict len:",len(hit_chunk_info_dict))
            if self.num_expand_context_chunk > 0:
                new_hit_chunk_info_dict={}
                new_reference_results = []
                for corpus_id, hit_chunk_info in hit_chunk_info_dict.items():
                    hit_chunk=hit_chunk_info["hit_chunk"]
                    is_rerank=hit_chunk_info["is_rerank"]
                    if not is_rerank and self.rerank_model is not None:
                        continue
                    expanded_reference = self.sim_model.corpus.get(corpus_id - 1, '') + hit_chunk
                    for i in range(self.num_expand_context_chunk):
                        expanded_reference += self.sim_model.corpus.get(corpus_id + i + 1, '')
                    new_reference_results.append(expanded_reference)
                    hit_chunk_info['hit_chunk']=expanded_reference
                    new_hit_chunk_info_dict[corpus_id]=hit_chunk_info
                reference_results = new_reference_results
                print(f"hit_chunk_info_dict len:{len(hit_chunk_info_dict)},new_reference_results:{len(new_reference_results)}")
        # import json 
        # logger.debug(json.dumps(new_hit_chunk_info_dict,indent=4,ensure_ascii=False))
        new_hit_chunk_info_dict={i+1:v for i,(k,v) in enumerate(new_hit_chunk_info_dict.items())}
        return reference_results,new_hit_chunk_info_dict

    def predict_stream(
            self,
            query: str,
            max_length: int = 512,
            context_len: int = 2048*20,
            temperature: float = 0.7,
    ):
        """Generate predictions stream."""
        stop_str = self.tokenizer.eos_token if self.tokenizer.eos_token else "</s>"
        if not self.enable_history:
            self.history = []
        if self.sim_model.corpus:
            reference_results,hit_chunk_info_dict = self.get_reference_results(query)
            # print(len(reference_results))
            if reference_results:
                reference_results = self._add_source_numbers(reference_results)
                context_str = '\n'.join(reference_results)[:(context_len - len(PROMPT_TEMPLATE))]
            else:
                context_str = ''
            prompt = PROMPT_TEMPLATE.format(context_str=context_str, query_str=query)
        else:
            prompt = query
        
        self.context_str=context_str
        # logger.debug(f"prompt: {prompt}")
        self.history.append([prompt, ''])
        response = ""
        if self.openai_api_key:
        # 使用 OpenAI 的流式回答
            for new_text in self.openai_stream_generate_answer(prompt, max_tokens=max_length, temperature=temperature):
                response += new_text
                yield response
        else:
            for new_text in self.stream_generate_answer(
                    max_new_tokens=max_length,
                    temperature=temperature,
                    context_len=context_len,
            ):
                if new_text != stop_str:
                    response += new_text
                    yield response

    def predict(
            self,
            query: str,
            max_length: int = 512,
            context_len: int = 2048*20,
            temperature: float = 0.7,
    ):
        """Query from corpus."""
        reference_results = []
        if not self.enable_history:
            self.history = []
        if self.sim_model.corpus:
            reference_results,hit_chunk_info_dict = self.get_reference_results(query)
            if reference_results:
                # print(reference_results)
                reference_results = self._add_source_numbers(reference_results)
                # print(f"after reference_results len:{len(reference_results)}")
                context_str = '\n'.join(reference_results)[:(context_len - len(PROMPT_TEMPLATE))]
            else:
                context_str = ''
            prompt = PROMPT_TEMPLATE.format(context_str=context_str, query_str=query)
        else:
            prompt = query
        # logger.debug(f"prompt: {prompt}")
        self.history.append([prompt, ''])
        response = ""
        if self.openai_api_key:
        # 使用 OpenAI 的流式回答
            for new_text in self.openai_stream_generate_answer(prompt, max_tokens=max_length, temperature=temperature):
                response += new_text
        else:
            for new_text in self.stream_generate_answer(
                    max_new_tokens=max_length,
                    temperature=temperature,
                    context_len=context_len,
            ):
                response += new_text
        response = response.strip()
        self.history[-1][1] = response
        return response, hit_chunk_info_dict,prompt

    def query(self, query: str, **kwargs):
        return self.predict(query, **kwargs)

    def save_corpus_emb(self):
        dir_name = self.get_file_hash(self.corpus_files)
        save_dir = os.path.join(self.save_corpus_emb_dir, dir_name)
        if hasattr(self.sim_model, 'save_corpus_embeddings'):
            self.sim_model.save_corpus_embeddings(save_dir)
            logger.debug(f"Saving corpus embeddings to {save_dir}")
        return save_dir

    def load_corpus_emb(self, emb_dir: str):
        if hasattr(self.sim_model, 'load_corpus_embeddings'):
            logger.debug(f"Loading corpus embeddings from {emb_dir}")
            self.sim_model.load_corpus_embeddings(emb_dir)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--sim_model_name", type=str, default="shibing624/text2vec-base-multilingual")
    parser.add_argument("--gen_model_type", type=str, default="auto")
    parser.add_argument("--gen_model_name", type=str, default="Qwen/Qwen2-0.5B-Instruct")
    parser.add_argument("--lora_model", type=str, default=None)
    parser.add_argument("--rerank_model_name", type=str, default=None)
    parser.add_argument("--corpus_files", type=str, default="data/sample.pdf")
    parser.add_argument("--device", type=str, default=None)
    parser.add_argument("--int4", action='store_true', help="use int4 quantization")
    parser.add_argument("--int8", action='store_true', help="use int8 quantization")
    parser.add_argument("--chunk_size", type=int, default=220)
    parser.add_argument("--chunk_overlap", type=int, default=0)
    parser.add_argument("--num_expand_context_chunk", type=int, default=1)
    args = parser.parse_args()
    print(args)
    sim_model = BertSimilarity(model_name_or_path=args.sim_model_name, device=args.device)
    m = Rag(
        # similarity_model=sim_model,
        generate_model_type=args.gen_model_type,
        generate_model_name_or_path=args.gen_model_name,
        lora_model_name_or_path=args.lora_model,
        device=args.device,
        int4=args.int4,
        int8=args.int8,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        corpus_files=args.corpus_files.split(','),
        num_expand_context_chunk=args.num_expand_context_chunk,
        rerank_model_name_or_path=args.rerank_model_name,
    )
    response, hit_chunk_info_dict,prompt = m.predict('介绍下康复师.')
    print(response)
